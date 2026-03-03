import os
import streamlit as st
import mammoth
import io
import re
import time
import base64
from docx import Document

# Load both jszip and docx-preview from local files.
# JSZip sets window.JSZip; docx-preview's UMD then falls through to the
# browser-global branch which reads window.JSZip and sets window.docxPreview.
_BASE = os.path.dirname(__file__)
with open(os.path.join(_BASE, "jszip.min.js"), "r", encoding="utf-8") as _f:
    _JSZIP_JS = _f.read()
with open(os.path.join(_BASE, "docx_preview.min.js"), "r", encoding="utf-8") as _f:
    _DOCX_PREVIEW_JS = _f.read()


st.set_page_config(layout="wide", page_title="RevShift.ai - Conga → OmniStudio", page_icon="🪄")

# ── Session State ──
for key, default in [
    ('file_bytes', None), ('filename', None),
    ('converted_bytes', None), ('stats', None),
    ('detected_fields', []),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ── Regex Pattern for Conga merge fields ──
CONGA_PATTERN = re.compile(r'\{![^}]+\}|\{\{.*?\}\}|<<.*?>>')

# ────────────────────────────────────────────
#  docx-preview: renders .docx faithfully
#  using the JS library loaded from CDN
# ────────────────────────────────────────────
def render_docx_preview(docx_bytes, container_id="docx-container", height=800):
    """Render a .docx file using the docx-preview JS library inside a Streamlit HTML component.
    This preserves colors, backgrounds, images, fonts, and table styling."""
    b64 = base64.b64encode(docx_bytes).decode('utf-8')
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
      <meta charset="utf-8">
      <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ background: #e8e8e8; }}
        #{container_id} {{
          background: #e8e8e8;
          padding: 10px 0;
          min-height: {height - 20}px;
        }}
        #{container_id} > section.docx {{
          background: white;
          margin: 10px auto;
          box-shadow: 0 2px 12px rgba(0,0,0,0.18);
          padding: 0;
        }}
      </style>
    </head>
    <body>
      <div id="{container_id}">
        <p style="text-align:center; padding: 40px; color: #888;">Rendering document…</p>
      </div>

      <!-- JSZip and docx-preview embedded inline — no CDN needed, works inside Streamlit iframe -->
      <script>{_JSZIP_JS}</script>
      <script>{_DOCX_PREVIEW_JS}</script>
      <script>
        (function() {{
          try {{
            var raw = atob("{b64}");
            var arr = new Uint8Array(raw.length);
            for (var i = 0; i < raw.length; i++) {{ arr[i] = raw.charCodeAt(i); }}

            var container = document.getElementById("{container_id}");

            // docx-preview exports globally to window.docx
            docx.renderAsync(arr.buffer, container, null, {{
              className: "docx",
              inWrapper: true,
              ignoreWidth: false,
              ignoreHeight: false,
              ignoreFonts: false,
              breakPages: true,
              ignoreLastRenderedPageBreak: false,
              experimental: true,
              trimXmlDeclaration: true,
              renderHeaders: true,
              renderFooters: true,
              renderFootnotes: true,
              renderEndnotes: true,
            }}).then(function() {{
              console.log("docx-preview: done");
            }}).catch(function(err) {{
              container.innerHTML = '<p style="color:red;padding:20px;">Error: ' + err.message + '</p>';
            }});
          }} catch(e) {{
            document.getElementById("{container_id}").innerHTML =
              '<p style="color:red;padding:20px;">Setup error: ' + e.message + '</p>';
          }}
        }})();
      </script>
    </body>
    </html>
    """
    st.components.v1.html(html, height=height, scrolling=True)

# ── Conversion helpers ──

def convert_match(m):
    """Convert a single Conga regex match to OmniStudio syntax."""
    text = m.group(0)
    inner = None
    if text.startswith('{!') and text.endswith('}'):
        inner = text[2:-1]
    elif text.startswith('{{') and text.endswith('}}'):
        inner = text[2:-2]
    elif text.startswith('<<') and text.endswith('>>'):
        inner = text[2:-2]
    if inner is None:
        return text
    if inner.startswith('#if '):
        return f"%#if {inner[4:].replace('.', ':')}%"
    elif inner.startswith('#'):
        return f"%#{inner[1:].replace('.', ':')}%"
    elif inner.startswith('/'):
        return f"%/{inner[1:].replace('.', ':')}%"
    else:
        return f"%{inner.replace('.', ':')}%"

def update_stats(stats, original, converted):
    stats['total'] += 1
    if converted.startswith('%#if '):
        stats['conditions'] += 1
    elif converted.startswith('%#') or converted.startswith('%/'):
        stats['loops'] += 1
    else:
        stats['converted'] += 1
    stats['fields'].append({'original': original, 'converted': converted})

def process_paragraph(p, stats):
    full_text = p.text
    matches = list(CONGA_PATTERN.finditer(full_text))
    if not matches:
        return
    run_match_count = sum(1 for r in p.runs for _ in CONGA_PATTERN.finditer(r.text))
    if run_match_count == len(matches):
        for r in p.runs:
            if CONGA_PATTERN.search(r.text):
                original_text = r.text
                new_text = original_text
                for m in CONGA_PATTERN.finditer(original_text):
                    converted = convert_match(m)
                    new_text = new_text.replace(m.group(0), converted)
                    update_stats(stats, m.group(0), converted)
                r.text = new_text
    else:
        original_text = full_text
        new_text = original_text
        for m in CONGA_PATTERN.finditer(original_text):
            converted = convert_match(m)
            new_text = new_text.replace(m.group(0), converted)
            update_stats(stats, m.group(0), converted)
        style = p.style
        p.clear()
        p.add_run(new_text)
        p.style = style

def process_docx(file_bytes):
    """Run conversion on all paragraphs and table cells."""
    doc = Document(io.BytesIO(file_bytes))
    stats = {'total': 0, 'converted': 0, 'unsupported': 0, 'loops': 0, 'conditions': 0, 'fields': []}
    for p in doc.paragraphs:
        process_paragraph(p, stats)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p, stats)
    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue(), stats

def detect_fields(docx_bytes):
    """Scan all paragraphs & table cells for Conga merge fields."""
    doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_text += "\n" + "\n".join(p.text for p in c.paragraphs)
    return CONGA_PATTERN.findall(all_text)

# ────────────────────────────────────────────
#  UI: Wizard Tabs + Progress
# ────────────────────────────────────────────

st.title("✨ RevShift.ai")
st.caption("Intelligently convert Conga `.docx` merge templates into OmniStudio Document Generation format.")

# Progress indicator
has_file = st.session_state['file_bytes'] is not None
has_converted = st.session_state['converted_bytes'] is not None
step = 1 + int(has_file) + int(has_converted)

st.markdown("<br>", unsafe_allow_html=True)
st.progress(min(step / 3, 1.0), text=f"Progress: Step {step} of 3")
st.markdown("<br>", unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📂 1. Upload", "📄 2. Original Preview", "⚙️ 3. Analyse & Convert",
    "✨ 4. Converted Preview", "📥 5. Download"
])

# ── STEP 1: Upload ──
with tab1:
    st.header("Upload Conga Template")
    uploaded_file = st.file_uploader("Choose a **.docx** file", type="docx")
    if uploaded_file is not None:
        st.session_state['file_bytes'] = uploaded_file.getvalue()
        st.session_state['filename'] = uploaded_file.name
        size_kb = len(st.session_state['file_bytes']) / 1024
        st.success(f"✅  **{uploaded_file.name}**  ({size_kb:.1f} KB) uploaded successfully.")
        st.info("👉 Proceed to the **2. Original Preview** tab or skip straight to **3. Analyse & Convert**.")

# ── STEP 2: Original Preview ──
with tab2:
    st.header("Original Document Preview")
    if st.session_state['file_bytes']:
        col_preview, col_sidebar = st.columns([3, 1])
        with col_preview:
            render_docx_preview(st.session_state['file_bytes'], container_id="orig-preview", height=800)
        with col_sidebar:
            st.subheader("Detected Merge Fields")
            fields = detect_fields(st.session_state['file_bytes'])
            st.metric("Total Fields", len(fields))
            for f in fields:
                st.code(f, language=None)
    else:
        st.info("⬆ Upload a file first in the **① Upload** tab.")

# ── STEP 3: Analyse & Convert ──
with tab3:
    st.header("Analyse & Convert Engine")
    if st.session_state['file_bytes']:
        st.markdown("Click below to analyze Conga tags and convert them into OmniStudio syntax.")
        if st.button("🚀 Run RevShift.ai Conversion", type="primary", use_container_width=True):
            t0 = time.time()
            with st.spinner("🤖 Analyzing syntax and reconstructing document..."):
                time.sleep(0.5) # Slight delay for UI feedback
                st.session_state['converted_bytes'], st.session_state['stats'] = process_docx(st.session_state['file_bytes'])
            elapsed = time.time() - t0
            st.session_state['stats']['time'] = elapsed
            st.success(f"🎉 Conversion completed successfully in **{elapsed:.2f}s**!")
            st.balloons()

        if st.session_state['stats']:
            s = st.session_state['stats']
            st.subheader("📊 Conversion Report")
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("Total Fields", s['total'])
            c2.metric("Basic Fields", s['converted'])
            c3.metric("Table Loops", s['loops'])
            c4.metric("Conditions", s['conditions'])
            c5.metric("⏱ Time", f"{s.get('time', 0):.2f}s")
            st.divider()
            st.subheader("Field Mapping Table")
            st.dataframe(s['fields'], use_container_width=True)
    else:
        st.info("⬆ Upload a file first.")

# ── STEP 4: Converted Preview (full width) ──
with tab4:
    st.header("Converted Document Preview")
    if st.session_state['converted_bytes']:
        st.info("Showing the fully converted document. OmniStudio syntax should now be visible.")
        render_docx_preview(st.session_state['converted_bytes'], container_id="cmp-conv", height=800)
    else:
        st.warning("⬅ Complete the **3. Analyse & Convert** step first to see the preview.")

# ── STEP 5: Download ──
with tab5:
    st.header("Download Converted Template")
    if st.session_state['converted_bytes']:
        base = st.session_state['filename'].rsplit('.', 1)[0]
        new_name = f"{base}_omnistudio_converted.docx"
        st.download_button(
            label="⬇  Download .docx",
            data=st.session_state['converted_bytes'],
            file_name=new_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        st.info(f"File will be saved as **{new_name}**")
    else:
        st.info("⬅ Complete the conversion step first.")
